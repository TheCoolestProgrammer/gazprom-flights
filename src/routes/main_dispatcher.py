from datetime import datetime
import os
from typing import Optional

from fastapi import APIRouter, Body, HTTPException, Request, Depends, Form, Response, status,UploadFile, File
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse
from src.models.airport import Airport
from src.models.pilot import Pilot
from src.crud.flight import create_flights_bulk
from src.crud.excel_generator import generate_both_excel_files
from src.database import SessionDep
from src.dependencies import get_current_user, RoleChecker
# from src.models.flight_route import FlightRoute
from src.models.user import User, Role
from src.models.passenger import Passenger, RequestStatus,Gender,TripPurpose, GTURelation
from src.models.department import Department
from src.models.flights import Flight
from src.schemas.flight import FlightCreate, FlightCreateForm, FlightResponse, FlightParseResponse, SelectedFlightsRequest
from src.parsers.docs_parser import parse_flight_docx
from docx import Document
from io import BytesIO
from src.templates_config import templates
import zipfile

router = APIRouter(prefix="/main_dispatcher", tags=["main_dispatcher"])


@router.get("/", response_class=HTMLResponse)
async def dashboard(
    request: Request,
    session: SessionDep,
    user: User = Depends(RoleChecker(Role.DISPATCHER_DIRECTOR))
    ):
    requests = session.query(Passenger).filter(Passenger.department_director_status==RequestStatus.CONFIRMED).order_by(Passenger.request_date.desc()).all()
    return templates.TemplateResponse(request=request, name="main_dispatcher/dashboard.html", context={
        "user": user,
        "requests": requests,
        "Status":RequestStatus
    })


@router.get("/edit/{passenger_id}", response_class=HTMLResponse)
async def edit_form(
    request: Request,
    passenger_id:int,
    session: SessionDep, 
    user: User = Depends(RoleChecker(Role.DISPATCHER_DIRECTOR))
    ):
    passenger = session.query(Passenger).filter(Passenger.id==passenger_id).first()
    departments = session.query(Department).all()
    # flight_routes = session.query(FlightRoute).all()
    airports = session.query(Airport).all()
    print(airports)
    return templates.TemplateResponse(request=request, name="main_dispatcher/edit.html", context={
        "user": user,
        "departments": departments,
        "passenger":passenger,
        "airports":airports,
        # "flight_routes":flight_routes
        "TripPurpose":TripPurpose,
        "GTURelation":GTURelation
    })

@router.post("/edit/{passenger_id}")
async def edit_request(
    request: Request,
    session: SessionDep,
    passenger_id:int,
    fullname: str = Form(...),
    passport: int = Form(...),
    flight_from: int = Form(...),
    birthdate: str = Form(...),
    gender: Gender = Form(...),
    trip_purpose: TripPurpose = Form(...),
    planning_date: str = Form(...),
    flight_to: int = Form(...),
    cargo_weight:float = Form(...),
    gtu_relation:GTURelation = Form(...),
    department_id:int = Form(...),
    notes:Optional[str] = Form(None),
    user: User = Depends(RoleChecker(Role.DISPATCHER_DIRECTOR))
):
    passenger = session.get(Passenger, passenger_id)
    
    passenger.fullname=fullname
    passenger.passport=passport
    passenger.flight_from_id=flight_from
    passenger.birthdate=birthdate # Нужна конвертация в date объект
    passenger.gender=gender
    passenger.trip_purpose=trip_purpose
    passenger.status=RequestStatus.PENDING
    # passenger.created_by=user.id
    passenger.planning_date=planning_date
    passenger.flight_to_id=flight_to
    passenger.cargo_weight=cargo_weight
    passenger.gtu_relation=gtu_relation
    passenger.department_id = department_id
    if notes:
        passenger.notes=notes

    session.commit()
    return RedirectResponse(url="/main_dispatcher/", status_code=303)

@router.patch("/change_status/{passenger_id}", response_class=HTMLResponse)
async def change_status(
    request: Request,
    session: SessionDep,
    passenger_id:int,
    request_status:RequestStatus = Body(..., embed=True),
    user: User = Depends(RoleChecker(Role.DISPATCHER_DIRECTOR))
    ):
    passenger = session.get(Passenger,passenger_id)
    
    if not passenger:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, 
            detail=f"Пассажир с ID {passenger_id} не найден"
        )
    # try:
    passenger.main_dispatcher_status = request_status
    session.commit()
    
    return JSONResponse(
        content={"message": "Successfully changed"}, 
        status_code=status.HTTP_200_OK
    )


@router.post("/upload-docx", response_model=FlightParseResponse)
async def upload_flights_docx(
    db: SessionDep,
    file: UploadFile = File(..., description="DOCX файл с заявкой на полёты")
):
    """
    Загружает DOCX файл с заявкой на выполнение полетов,
    парсит его и сохраняет данные в базу данных.
    """
    # Проверка типа файла
    if not file.filename.endswith('.docx'):
        raise HTTPException(
            status_code=400,
            detail="Неверный формат файла. Ожидается файл с расширением .docx"
        )
    
    try:
        # Чтение содержимого файла
        file_content = await file.read()
        
        # Парсинг документа
        parsed_data = parse_flight_docx(file_content)
        
        if not parsed_data["departure_date"]:
            raise HTTPException(
                status_code=400,
                detail="Не удалось найти дату выполнения полётов в документе"
            )
        
        if not parsed_data["flights"]:
            raise HTTPException(
                status_code=400,
                detail="Не удалось найти рейсы в документе"
            )
        
        # Подготовка данных для сохранения
        flights_to_create = []
        for flight_data in parsed_data["flights"]:
            flight_create = FlightCreate(
                aircraft_type=flight_data["aircraft_type"],
                flight_number=flight_data["flight_number"],
                departure_date=parsed_data["departure_date"],
                departure_time=flight_data["departure_time"],
                place_number=flight_data["place_number"],
                route=flight_data["route"]
            )
            flights_to_create.append(flight_create)
        
        # Сохранение в БД
        saved_flights = create_flights_bulk(db, flights_to_create)
        
        # Преобразование в response schema
        flights_response = [
            FlightResponse(
                id=flight.id,
                aircraft_type=flight.aircraft_type,
                flight_number=flight.flight_number,
                departure_date=flight.departure_date,
                departure_time=flight.departure_time,
                place_number=flight.place_number,
                route=flight.route
            )
            for flight in saved_flights
        ]
        
        return FlightParseResponse(
            status="success",
            message=f"Успешно обработано {len(saved_flights)} рейсов",
            flights_parsed=len(saved_flights),
            flights_saved=flights_response
        )
        
    except HTTPException:
        raise
    except Exception as e:
        # logger.error(f"Ошибка при обработке файла: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Внутренняя ошибка сервера при обработке файла: {str(e)}"
        )

@router.get("/flights", response_class=HTMLResponse)
async def flights_page(request: Request, db: SessionDep):
    flights = db.query(Flight).order_by(Flight.departure_date.desc()).all()
    pilots = db.query(Pilot).all()
    return templates.TemplateResponse(request=request, name="main_dispatcher/flights.html", context={
        "flights": flights,
        "pilots": pilots
    })

@router.post("/create-from-form", response_model=dict)
async def create_flight_from_form(
    flight_data: FlightCreateForm,  # Используем Pydantic модель
    db: SessionDep
):
    """
    Создаёт рейс из формы, сохраняет в БД и генерирует DOCX-файл.
    """
    try:
        # Логируем полученные данные для отладки
        print(f"Получены данные: {flight_data}")
        
        # Проверяем, нет ли уже рейса с таким же номером на эту дату
        existing_flight = db.query(Flight).filter(
            Flight.flight_number == flight_data.flight_number,
            Flight.departure_date == flight_data.departure_date
        ).first()
        
        if existing_flight:
            raise HTTPException(
                status_code=400, 
                detail=f"Рейс с номером {flight_data.flight_number} на дату {flight_data.departure_date} уже существует"
            )
        
        # Сохраняем в БД
        new_flight = Flight(
            aircraft_type=flight_data.aircraft_type,
            flight_number=flight_data.flight_number,
            departure_date=flight_data.departure_date,
            departure_time=flight_data.departure_time,
            place_number=flight_data.place_number,
            route=flight_data.route,
            pilot_id=flight_data.pilot_id
        )
        db.add(new_flight)
        db.commit()
        db.refresh(new_flight)
        
        # Генерируем DOCX
        docx_bytes = generate_flight_docx(new_flight, flight_data.gzp)
        
        # Сохраняем файл
        os.makedirs("generated_docx", exist_ok=True)
        filename = f"flight_{new_flight.id}_{new_flight.departure_date}.docx"
        filepath = os.path.join("generated_docx", filename)
        with open(filepath, "wb") as f:
            f.write(docx_bytes)
        
        return {"status": "success", "flight_id": new_flight.id, "file": filename}
    
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        print(f"Ошибка: {str(e)}")  # Логируем ошибку
        raise HTTPException(status_code=500, detail=f"Ошибка при создании рейса: {str(e)}")

@router.get("/download/{flight_id}")
async def download_flight_docx(flight_id: int, db: SessionDep):
    """
    Скачивает сгенерированный DOCX-файл для рейса.
    """
    flight = db.query(Flight).filter(Flight.id == flight_id).first()
    if not flight:
        raise HTTPException(status_code=404, detail="Рейс не найден")
    
    # Генерируем заново или берём сохранённый
    docx_bytes = generate_flight_docx(flight, "")
    
    # Отдаём файл
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=flight_{flight_id}.docx"}
    )

@router.get("/list", response_model=list[FlightResponse])
async def list_flights(db: SessionDep):
    """Список всех рейсов (для отображения в таблице)"""
    flights = db.query(Flight).order_by(Flight.departure_date.desc(), Flight.departure_time).all()
    return flights

def generate_flight_docx(flight: Flight, gzp: str) -> bytes:
    """
    Генерирует DOCX-документ в формате заявки на выполнение полётов.
    """
    doc = Document()
    
    # Заголовок
    title = doc.add_paragraph("Заявка на выполнение полетов")
    title.runs[0].bold = True
    
    doc.add_paragraph()  # Пустая строка
    
    # Дата (преобразуем в русский формат с днём недели)
    weekday_ru = {
        0: "понедельник", 1: "вторник", 2: "среда", 3: "четверг",
        4: "пятница", 5: "суббота", 6: "воскресенье"
    }
    weekday_name = weekday_ru[flight.departure_date.weekday()]
    date_str = f"{flight.departure_date.day:02d}.{flight.departure_date.month:02d}.{flight.departure_date.year} {weekday_name}"
    doc.add_paragraph(date_str)
    
    doc.add_paragraph()  # Пустая строка
    
    # Рейс (номер рейса используем как есть)
    flight_line = (f"1. {flight.aircraft_type} {flight.flight_number} ГЗП {gzp} "
                   f"время вылета {flight.departure_time.strftime('%H:%M')} "
                   f"кол-во кресел {flight.place_number}")
    doc.add_paragraph(flight_line)
    
    # Маршрут
    doc.add_paragraph(f"Маршрут: {flight.route}")
    
    # Сохраняем в байты
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


@router.get("/download-selected")
async def download_selected_flights(
    request: SelectedFlightsRequest,
    db: SessionDep
):
    """
    Скачивает выбранные рейсы в одном DOCX-файле.
    """
    if not request.flight_ids:
        raise HTTPException(status_code=400, detail="Не выбрано ни одного рейса")
    
    # Получаем рейсы из БД
    flights = db.query(Flight).filter(Flight.id.in_(request.flight_ids)).order_by(Flight.departure_date, Flight.departure_time).all()
    
    if not flights:
        raise HTTPException(status_code=404, detail="Рейсы не найдены")
    
    # Генерируем DOCX со всеми рейсами
    docx_bytes = generate_multiple_flights_docx(flights)
    
    # Генерируем имя файла с датами
    dates = set(f.departure_date for f in flights)
    if len(dates) == 1:
        filename = f"flights_{list(dates)[0]}.docx"
    else:
        filename = f"flights_{min(dates)}_{max(dates)}.docx"
    
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@router.get("/download-excel-tables/{flight_id}")
async def download_excel_tables(flight_id: int, db: SessionDep):
    """
    Скачивает два Excel файла с таблицами пассажиров для рейса:
    1. ОБРАЗЕЦ!!Список пассажиров.xlsx - манифест пассажиров
    2. Список_пассажиров_для_оформления_авиабилетов.xlsx - список для оформления билетов
    """
    flight = db.query(Flight).filter(Flight.id == flight_id).first()
    if not flight:
        raise HTTPException(status_code=404, detail="Рейс не найден")
    
    # Генерируем оба Excel файла
    manifest_bytes, ticket_list_bytes = generate_both_excel_files(flight, db)
    
    # Создаём ZIP архив с двумя файлами
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # Добавляем манифест
        manifest_filename = f"Манифест_рейс_{flight.flight_number}_{flight.departure_date}.xlsx"
        zip_file.writestr(manifest_filename, manifest_bytes)
        
        # Добавляем список для оформления билетов
        ticket_filename = f"Билеты_рейс_{flight.flight_number}_{flight.departure_date}.xlsx"
        zip_file.writestr(ticket_filename, ticket_list_bytes)
    
    zip_buffer.seek(0)
    
    return Response(
        content=zip_buffer.getvalue(),
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename=flight_{flight.flight_number}_{flight.departure_date}.zip"}
    )

def generate_multiple_flights_docx(flights: list[Flight]) -> bytes:
    """
    Генерирует DOCX-документ с несколькими рейсами.
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Заголовок
    title = doc.add_paragraph("Заявка на выполнение полетов")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)
    
    doc.add_paragraph()  # Пустая строка
    
    # Группируем рейсы по датам
    from collections import defaultdict
    flights_by_date = defaultdict(list)
    for flight in flights:
        flights_by_date[flight.departure_date].append(flight)
    
    # Для каждой даты создаём раздел
    for i, (date, date_flights) in enumerate(sorted(flights_by_date.items())):
        if i > 0:
            doc.add_page_break()
        
        # Преобразуем дату в русский формат с днём недели
        weekday_ru = {
            0: "понедельник", 1: "вторник", 2: "среда", 3: "четверг",
            4: "пятница", 5: "суббота", 6: "воскресенье"
        }
        weekday_name = weekday_ru[date.weekday()]
        date_str = f"{date.day:02d}.{date.month:02d}.{date.year} {weekday_name}"
        
        date_paragraph = doc.add_paragraph(date_str)
        date_paragraph.runs[0].bold = True
        date_paragraph.runs[0].font.size = Pt(14)
        
        doc.add_paragraph()  # Пустая строка
        
        # Добавляем каждый рейс
        for idx, flight in enumerate(date_flights, 1):
            # Получаем ГЗП (если нет в модели, можно добавить или передавать отдельно)
            # Для этого нужно добавить поле gzp в модель Flight
            # gzp = getattr(flight, 'gzp', '_____')
            
            flight_line = (f"{idx}. {flight.aircraft_type} {flight.flight_number} ГЗП "
                          f"время вылета {flight.departure_time.strftime('%H:%M')} "
                          f"кол-во кресел {flight.place_number}")
            doc.add_paragraph(flight_line)
            
            # Маршрут
            doc.add_paragraph(f"Маршрут: {flight.route}")
            
            # Добавляем пустую строку между рейсами, если не последний
            if idx < len(date_flights):
                doc.add_paragraph()
    
    # Сохраняем в байты
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

